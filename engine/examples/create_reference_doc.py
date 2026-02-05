#!/usr/bin/env python3
"""
Create a custom reference document for Pandoc DOCX export.

This document defines the styles that Pandoc will use when converting
markdown to DOCX format. Styles include:
- Title, Subtitle, Author, Date
- Heading 1, 2, 3 (APA 7th edition compatible)
- Normal text (double-spaced, Times New Roman 12pt)
- Block quotes
- List styles
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_reference_document():
    """Create a professionally styled reference document for Pandoc."""
    doc = Document()

    # Set document margins (1 inch all around - academic standard)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configure Normal style (base for all text)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    # Configure Title style
    title = doc.styles['Title']
    title.font.name = 'Times New Roman'
    title.font.size = Pt(16)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    title.paragraph_format.first_line_indent = Inches(0)

    # Configure Subtitle style
    subtitle = doc.styles['Subtitle']
    subtitle.font.name = 'Times New Roman'
    subtitle.font.size = Pt(14)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0, 0, 0)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.first_line_indent = Inches(0)

    # Configure Heading 1 (Major sections - centered, bold)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.page_break_before = False

    # Configure Heading 2 (Subsections - left-aligned, bold)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h2.paragraph_format.first_line_indent = Inches(0)

    # Configure Heading 3 (Sub-subsections - left-aligned, bold italic)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h3.paragraph_format.first_line_indent = Inches(0)

    # Configure Block Text / Quote style
    try:
        block_text = doc.styles['Block Text']
    except KeyError:
        block_text = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    block_text.font.name = 'Times New Roman'
    block_text.font.size = Pt(12)
    block_text.paragraph_format.left_indent = Inches(0.5)
    block_text.paragraph_format.right_indent = Inches(0.5)
    block_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    block_text.paragraph_format.first_line_indent = Inches(0)

    # Configure List Bullet style
    list_bullet = doc.styles['List Bullet']
    list_bullet.font.name = 'Times New Roman'
    list_bullet.font.size = Pt(12)
    list_bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(0)

    # Configure List Number style
    list_number = doc.styles['List Number']
    list_number.font.name = 'Times New Roman'
    list_number.font.size = Pt(12)
    list_number.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    list_number.paragraph_format.space_before = Pt(0)
    list_number.paragraph_format.space_after = Pt(0)

    # Add sample content to demonstrate styles (Pandoc needs at least one use of each style)
    doc.add_paragraph('Title', style='Title')
    doc.add_paragraph('Subtitle', style='Subtitle')
    doc.add_paragraph('Heading 1', style='Heading 1')
    doc.add_paragraph('Heading 2', style='Heading 2')
    doc.add_paragraph('Heading 3', style='Heading 3')
    doc.add_paragraph('Normal paragraph text with proper academic formatting.', style='Normal')
    doc.add_paragraph('Block quote text.', style='Block Text')
    doc.add_paragraph('Bullet item', style='List Bullet')
    doc.add_paragraph('Numbered item', style='List Number')

    # Save the reference document
    output_path = Path(__file__).parent / 'custom-reference.docx'
    doc.save(str(output_path))
    print(f"Created reference document: {output_path}")
    return output_path


if __name__ == '__main__':
    create_reference_document()
